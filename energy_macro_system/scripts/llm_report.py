#!/usr/bin/env python3
"""
LLM-based 1-year macroeconomic forecast report for Energy–Macro Shock Propagation Simulator.
Calls local Ollama, then writes a professional Word (.docx) report.
Run from project root or energy_macro_system. Requires: pip install python-docx
"""

import json
import os
import urllib.error
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Optional

# Optional: python-docx for Word output
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

OLLAMA_URL = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434").rstrip("/")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "gemma3:12b")
GENERATE_ENDPOINT = f"{OLLAMA_URL}/api/generate"
DEFAULT_OUTPUT_PATH = "1_year_macro_forecast.docx"
REQUEST_TIMEOUT = 120


def call_ollama(
    prompt: str,
    model: str = OLLAMA_MODEL,
    temperature: float = 0.2,
    timeout: int = REQUEST_TIMEOUT,
) -> str:
    """
    Call local Ollama /api/generate. Returns generated text.
    Raises on connection error or non-200; returns empty string if response has no text.
    """
    body = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": temperature},
    }
    data = json.dumps(body).encode("utf-8")
    req = urllib.request.Request(
        GENERATE_ENDPOINT,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            if resp.status != 200:
                raise RuntimeError(f"Ollama returned HTTP {resp.status}")
            out = json.loads(resp.read().decode("utf-8"))
    except urllib.error.URLError as e:
        raise RuntimeError(f"Ollama request failed: {e}") from e
    except json.JSONDecodeError as e:
        raise RuntimeError(f"Invalid JSON from Ollama: {e}") from e

    text = (out.get("response") or "").strip()
    return text


def _build_prompt(metrics: Dict[str, Any]) -> str:
    """Build a concise, structured prompt for 1-year outlook. No invented numbers."""
    lines = [
        "You are a macroeconomic analyst. Using ONLY the following model outputs, write a 1-year macroeconomic forecast report.",
        "Use plain English. Do not invent any numbers; only reference the figures below.",
        "Keep the total length under 300 words.",
        "",
        "Model outputs:",
    ]
    if "spectral_radius" in metrics:
        lines.append(f"- Spectral radius (stability): {metrics['spectral_radius']:.4f}")
    if "shock_amplification" in metrics:
        lines.append(f"- Shock amplification factor: {metrics['shock_amplification']:.4f}")
    if "recession_prob_12m" in metrics:
        lines.append(f"- 12-month recession probability (GDP < 0): {metrics['recession_prob_12m']:.2%}")
    if "inflation_prob_12m" in metrics:
        lines.append(f"- 12-month inflation stress probability: {metrics['inflation_prob_12m']:.2%}")
    if "gdp_mean_12m" in metrics:
        lines.append(f"- Mean GDP growth at 12 months: {metrics['gdp_mean_12m']:.2f}%")
    if "demand_stress" in metrics:
        lines.append(f"- Demand stress metric: {metrics['demand_stress']:.4f}")
    if "latest_values" in metrics and isinstance(metrics["latest_values"], dict):
        lines.append("- Latest values (selected series):")
        for k, v in metrics["latest_values"].items():
            if isinstance(v, (int, float)):
                lines.append(f"  - {k}: {v:.4f}")
            else:
                lines.append(f"  - {k}: {v}")
    lines.extend([
        "",
        "Provide:",
        "1. One short paragraph on overall stability and risk (spectral radius, shock amplification).",
        "2. One short paragraph on 1-year outlook: recession and inflation risks, using the probabilities above.",
        "3. One brief bullet list of key takeaways (no new numbers).",
        "Do not exceed 300 words. Do not invent data.",
    ])
    return "\n".join(lines)


def _narrative_to_docx(
    narrative: str,
    output_path: str,
    title: str = "1-Year Macroeconomic Forecast Report",
    date_generated: Optional[datetime] = None,
) -> None:
    """Write narrative into a Word document with title, date, and formatted body."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")

    doc = Document()
    date_generated = date_generated or datetime.now()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    # Date
    doc.add_paragraph()
    p = doc.add_paragraph(f"Generated: {date_generated.strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Body: split into paragraphs and bullets
    for block in narrative.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Bullet-like lines (leading "- " or "* ")
        if block.startswith("- ") or block.startswith("* "):
            for line in block.split("\n"):
                line = line.strip()
                if line.startswith("- "):
                    line = line[2:]
                elif line.startswith("* "):
                    line = line[2:]
                if line:
                    doc.add_paragraph(line, style="List Bullet")
            continue
        # Numbered or heading-like (e.g. "1. ...")
        if block.startswith("1.") or block.startswith("2.") or block.startswith("3."):
            first = block.split(".", 1)
            if len(first) == 2 and first[0].strip() in ("1", "2", "3"):
                doc.add_paragraph(first[1].strip())
                continue
        doc.add_paragraph(block)

    doc.save(output_path)


def generate_forecast_report(
    metrics_dict: Dict[str, Any],
    output_path: str = DEFAULT_OUTPUT_PATH,
    model: str = OLLAMA_MODEL,
    temperature: float = 0.2,
) -> str:
    """
    Build prompt from metrics_dict, call Ollama, write Word report.
    Returns the generated narrative text. Raises on Ollama or file errors.
    """
    prompt = _build_prompt(metrics_dict)
    narrative = call_ollama(prompt, model=model, temperature=temperature)
    if not narrative:
        raise RuntimeError("Ollama returned empty response")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    _narrative_to_docx(narrative, str(out), date_generated=datetime.now())
    return narrative


def main() -> None:
    """Example: generate report from sample metrics and save to 1_year_macro_forecast.docx."""
    # Example metrics (replace with real outputs from your simulator)
    metrics = {
        "spectral_radius": 0.92,
        "shock_amplification": 1.15,
        "recession_prob_12m": 0.18,
        "inflation_prob_12m": 0.22,
        "gdp_mean_12m": 1.8,
        "demand_stress": 0.05,
        "latest_values": {
            "oil_return": 0.01,
            "inflation_change": 0.3,
            "gdp_growth": 2.1,
        },
    }

    if not HAS_DOCX:
        print("Install python-docx: pip install python-docx")
        return

    try:
        narrative = generate_forecast_report(metrics, output_path=DEFAULT_OUTPUT_PATH)
        print("Report generated successfully.")
        print(f"Narrative length: {len(narrative)} chars")
        print(f"Saved to: {Path(DEFAULT_OUTPUT_PATH).resolve()}")
    except RuntimeError as e:
        print(f"Error: {e}")
        raise


if __name__ == "__main__":
    main()
